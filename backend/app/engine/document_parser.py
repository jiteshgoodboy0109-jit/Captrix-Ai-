import io
import re
import math
import json
import zipfile
import xml.etree.ElementTree as ET
import numpy as np
import pandas as pd
from typing import Dict, List, Any, Tuple, Optional
from app.engine.currency_engine import identify_currency

ACCOUNT_TYPE_RULES = {
    "REVENUE": [r"revenue", r"sales", r"income", r"turnover", r"fees earned", r"service revenue", r"gain"],
    "EXPENSE": [r"expense", r"cogs", r"cost of goods", r"salary", r"wages", r"rent", r"utility", r"depreciation", r"tax", r"interest", r"supplies", r"advertising", r"freight", r"payroll"],
    "ASSET": [r"asset", r"cash", r"bank", r"receivable", r"debtor", r"inventory", r"stock", r"prepaid", r"building", r"equipment", r"machinery", r"vehicle", r"investment", r"land"],
    "LIABILITY": [r"payable", r"creditor", r"debt", r"loan", r"borrowing", r"liability", r"accrued", r"mortgage", r"overdraft", r"tax payable", r"gst payable"],
    "EQUITY": [r"capital", r"equity", r"retained earnings", r"common stock", r"share capital", r"drawings", r"reserves"]
}

COMPILED_ACCOUNT_RULES = {
    category: [re.compile(p, re.IGNORECASE) for p in patterns]
    for category, patterns in ACCOUNT_TYPE_RULES.items()
}

def clean_value(val: Any) -> float:
    """Safely convert cell value to float. Preserves signs and returns 0.0 for missing/invalid cells."""
    v = clean_value_or_none(val)
    return v if v is not None else 0.0

def clean_value_or_none(val: Any) -> Optional[float]:
    """Safely convert cell value to float. Returns None for missing/blank/unreported cells."""
    if is_blank_value(val):
        return None
    if isinstance(val, (int, float)):
        v = float(val)
        if math.isnan(v) or math.isinf(v):
            return None
        return v
    s = str(val).replace("$", "").replace("₹", "").replace("€", "").replace("£", "").replace(",", "").strip()
    if s.startswith("(") and s.endswith(")"):
        s = "-" + s[1:-1]
    try:
        v = float(s)
        if math.isnan(v) or math.isinf(v):
            return None
        return v
    except ValueError:
        return None

def is_blank_value(val: Any) -> bool:
    """Check if a cell is truly blank/unreported vs zero."""
    if pd.isna(val) or val is None:
        return True
    s = str(val).strip()
    if s == "" or s.lower() in ["nan", "none", "null", "-", "--", "n/a", "not reported"]:
        return True
    return False

def sanitize_json_data(obj: Any) -> Any:
    """Recursively replace any NaN or Infinity floats with 0.0 so JSON serialization never fails."""
    if isinstance(obj, dict):
        return {k: sanitize_json_data(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [sanitize_json_data(v) for v in obj]
    elif isinstance(obj, float):
        if math.isnan(obj) or math.isinf(obj):
            return 0.0
        return obj
    return obj

def is_summary_or_total_row(name: str) -> bool:
    """Check if line item is an aggregated summary or total row to prevent double-counting."""
    name_lower = name.lower().strip()
    exact_totals = [
        "total", "subtotal", "sub-total", "grand total", "total assets", "total current assets",
        "total non-current assets", "total liabilities", "total current liabilities",
        "total long-term liabilities", "total long term liabilities", "total non-current liabilities",
        "total equity", "total owner's equity", "total liabilities and equity", "total liabilities & equity",
        "total revenue", "total income", "total sales", "total expenses", "total opex",
        "gross profit", "operating profit", "ebitda", "ebit", "pbt", "profit before tax",
        "net profit", "net income", "profit for the year", "net profit after tax", "pat"
    ]
    if name_lower in exact_totals:
        return True
    if name_lower.startswith("total ") or name_lower.startswith("subtotal ") or name_lower.startswith("sub-total "):
        return True
    if name_lower.endswith(" total") or name_lower.endswith(" subtotal"):
        return True
    return False

def classify_account(name: str, sheet_context: str = "") -> str:
    name_lower = name.lower()
    ctx_lower = sheet_context.lower()
    
    # Specific account mapping overrides for Non-Financial / Metrics & Cash Flow
    if any(kw in name_lower for kw in ["cash from", "operating activity", "investing activity", "financing activity", "net cash flow"]):
        return "CASH_FLOW"
    if any(kw in name_lower for kw in ["asset", "assets", "block", "ppe", "property, plant", "property plant"]):
        if "deferred tax asset" not in name_lower and "tax asset" not in name_lower:
            return "ASSET"
    if "liability" in name_lower or "liabilities" in name_lower:
        return "PAYABLE_LIABILITY" if "payable" in name_lower else "LIABILITY"
    metric_keywords = [
        "no. of equity shares", "equity shares in cr", "number of shares", "share count",
        "shares outstanding", "face value", "eps", "earnings per share", "diluted eps",
        "basic eps", "book value per share", "dividend per share", "number of employees",
        "pe ratio", "p/e", "margin %", "growth %", "yield", "price to earnings",
        "bonus shares", "new bonus shares", "bonus share", "bonus share count"
    ]
    if any(kw in name_lower for kw in metric_keywords) or ("shares" in name_lower and "capital" not in name_lower and "equity" not in name_lower):
        return "METRIC"
    if "capital work" in name_lower or "work in progress" in name_lower or "cwip" in name_lower:
        return "ASSET"
    if "cash" in name_lower or "bank" in name_lower:
        return "CASH_ASSET"
    if "receivable" in name_lower or "debtor" in name_lower:
        return "RECEIVABLE_ASSET"
    if "change in inventory" in name_lower or "changes in inventory" in name_lower or "inventory change" in name_lower:
        return "COGS"
    if "inventory" in name_lower or "stock" in name_lower:
        return "INVENTORY_ASSET"
    if "payable" in name_lower or "creditor" in name_lower:
        return "PAYABLE_LIABILITY"
    if "borrowing" in name_lower or "debt" in name_lower or "loan" in name_lower:
        return "DEBT_LIABILITY"
    if "interest" in name_lower or "finance cost" in name_lower or "finance charge" in name_lower or "borrowing cost" in name_lower:
        return "INTEREST_EXPENSE"
    if "depreciation" in name_lower or "amortisation" in name_lower or "amortization" in name_lower or "depr" in name_lower:
        return "DEPRECIATION_EXPENSE"
    if "tax" in name_lower or "taxes" in name_lower or "income tax" in name_lower or "provision for tax" in name_lower:
        if "payable" not in name_lower and "asset" not in name_lower and "deferred tax asset" not in name_lower and "deferred tax liability" not in name_lower:
            return "TAX_EXPENSE"

    if any(kw in name_lower for kw in ["dividend", "price:", "price", "report date"]):
        return "METRIC"
    if "net profit" in name_lower or "net income" in name_lower or "profit after tax" in name_lower:
        return "NET_INCOME"
    if "operating profit" in name_lower or "pbt" in name_lower or "profit before tax" in name_lower:
        return "OPERATING_INCOME"
    if any(kw in name_lower for kw in ["asset", "assets"]):
        if "deferred tax asset" not in name_lower and "tax asset" not in name_lower:
            return "ASSET"
    if any(kw in name_lower for kw in ["exp", "expense", "selling", "admin", "mfr", "overhead", "cost"]):
        return "EXPENSE"

    for acct_type, patterns in ACCOUNT_TYPE_RULES.items():
        for p in patterns:
            if re.search(p, name_lower):
                return acct_type

    # Fallback to sheet context
    if "income" in ctx_lower or "sales" in ctx_lower or "profit" in ctx_lower:
        return "REVENUE" if ("cost" not in name_lower and "expense" not in name_lower and "exp" not in name_lower) else "EXPENSE"
    if "balance" in ctx_lower:
        asset_keywords = ["asset", "block", "cwip", "investment", "building", "property", "equipment", "machinery", "goodwill", "receivable", "cash", "bank", "inventory"]
        if any(ak in name_lower for ak in asset_keywords):
            return "ASSET"
        if any(lk in name_lower for lk in ["payable", "liability", "liabilities", "debt", "borrowing", "provision"]):
            return "LIABILITY"
        if any(ek in name_lower for ek in ["equity", "share", "capital", "reserve", "surplus"]):
            return "EQUITY"

    return "EXPENSE" if ("cost" in name_lower or "fee" in name_lower or "exp" in name_lower) else "EXPENSE"

def detect_company_and_currency(sheet_data: Dict[str, pd.DataFrame], filename: str) -> Dict[str, Any]:
    """Inspect top 15 rows and columns of all sheets to detect exact Company Name, Currency, and Unit Scale."""
    detected_company = ""
    detected_currency = "USD"
    detected_unit = "Units"

    company_candidates = []
    
    # Currency symbols & patterns
    currency_map = {
        "₹": "INR", "inr": "INR", "rupees": "INR", "rs": "INR", "rs.": "INR",
        "$": "USD", "usd": "USD", "dollar": "USD",
        "€": "EUR", "eur": "EUR", "euro": "EUR",
        "£": "GBP", "gbp": "GBP", "pound": "GBP",
        "cad": "CAD", "aud": "AUD", "jpy": "JPY"
    }
    
    unit_map = [
        ("crore", "₹ Crores"), ("crores", "₹ Crores"), ("cr", "₹ Crores"), ("cr.", "₹ Crores"),
        ("lakh", "Lakhs"), ("lakhs", "Lakhs"), ("lac", "Lakhs"), ("lacs", "Lakhs"),
        ("million", "$ Millions"), ("millions", "$ Millions"), ("mn", "Millions"), ("m", "Millions"),
        ("thousand", "Thousands"), ("thousands", "Thousands"), ("k", "Thousands"),
        ("billion", "Billions"), ("billions", "Billions"), ("bn", "Billions")
    ]

    for sheet_name, df in sheet_data.items():
        if df.empty:
            continue
        
        # Scan columns first
        for col in df.columns:
            col_str = str(col).strip()
            col_lower = col_str.lower()
            
            iso_found, _ = identify_currency(col_str, filename)
            if iso_found != "USD" or detected_currency == "USD":
                detected_currency = iso_found
            
            for kw, u in unit_map:
                if kw in col_lower:
                    detected_unit = u
                    break
                    
            m_label = re.search(r'(?:company|entity|corporate|organization)\s*(?:name)?\s*[:\-]\s*([A-Za-z0-9\s.,]+)', col_str, re.IGNORECASE)
            if m_label:
                cand = m_label.group(1).strip()
                if cand and len(cand) < 80:
                    company_candidates.append(cand)
            
            if any(suffix in col_str.upper() for suffix in [" LTD", " LIMITED", " INC", " CORP", " CORPORATION", " LLC", " GROUP", " HOLDINGS", " PLC", " CO."]):
                cleaned = re.sub(r'[\(\)\[\]]', '', col_str).strip()
                if len(cleaned) < 80 and not any(kw in cleaned.lower() for kw in ["statement", "balance sheet", "income statement", "profit & loss"]):
                    company_candidates.append(cleaned)
        
        # Scan top 15 rows
        top_slice = df.iloc[:15]
        for _, row in top_slice.iterrows():
            for cell in row.values:
                if pd.isna(cell):
                    continue
                cell_str = str(cell).strip()
                cell_lower = cell_str.lower()
                
                # Check currency
                for kw, curr in currency_map.items():
                    if kw in cell_lower and detected_currency == "USD":
                        detected_currency = curr
                        
                # Check unit scale
                for kw, u in unit_map:
                    if kw in cell_lower:
                        detected_unit = u
                        break
                        
                # Check Company Name candidates
                # 1. Look for explicit labels
                m_label = re.search(r'(?:company|entity|corporate|organization)\s*(?:name)?\s*[:\-]\s*([A-Za-z0-9\s.,]+)', cell_str, re.IGNORECASE)
                if m_label:
                    cand = m_label.group(1).strip()
                    if cand and len(cand) < 80:
                        company_candidates.append(cand)
                
                # 2. Look for suffixes
                if any(suffix in cell_str.upper() for suffix in [" LTD", " LIMITED", " INC", " CORP", " CORPORATION", " LLC", " GROUP", " HOLDINGS", " PLC", " CO."]):
                    # Clean title line
                    cleaned = re.sub(r'[\(\)\[\]]', '', cell_str).strip()
                    if len(cleaned) < 80 and not any(kw in cleaned.lower() for kw in ["statement", "balance sheet", "income statement", "profit & loss"]):
                        company_candidates.append(cleaned)

    if company_candidates:
        detected_company = company_candidates[0]
    else:
        # Fallback to filename cleaning if workbook title is unstated
        base = re.sub(r'^\d+[_\-\s]*', '', filename) # remove leading numerical prefixes like 5_Wipro -> Wipro
        base = os.path.splitext(base)[0]
        cleaned = base.replace("_", " ").replace("-", " ")
        blacklist = ["financials", "financial", "statement", "statements", "tb", "ledger", "2024", "2025", "2026", "v1", "v2", "final", "excel", "sheet", "pdf", "docx", "doc", "csv", "txt", "json", "report", "accounts"]
        words = [w.capitalize() for w in cleaned.split() if w.lower() not in blacklist]
        if words:
            detected_company = " ".join(words)
        else:
            detected_company = cleaned.title() or "Enterprise Entity"

    return {
        "company_name": detected_company,
        "currency": detected_currency,
        "unit": detected_unit
    }

def detect_year_columns(headers: List[Any], top_rows: List[List[Any]]) -> Dict[int, str]:
    """Scan headers and provided top rows to locate financial year column indices."""
    year_map = {}
    
    # Combined search tokens across headers and top rows
    scan_rows = [headers] + top_rows
    
    for r_idx, row in enumerate(scan_rows):
        for col_idx, cell in enumerate(row):
            if col_idx in year_map or pd.isna(cell):
                continue
            cell_str = str(cell).strip()
            
            # Match FY2024, FY24, 2024, 31-Mar-2024, 31/03/2025, Q1 2025, 2026-03-31 00:00:00
            m_year = re.search(r'\b(201[5-9]|202[0-9]|2030)\b', cell_str)
            m_fy = re.search(r'\bFY\s*([0-9]{2,4})\b', cell_str, re.IGNORECASE)
            
            if m_year:
                year_map[col_idx] = m_year.group(1)
            elif m_fy:
                fy_val = m_fy.group(1)
                year_map[col_idx] = f"20{fy_val}" if len(fy_val) == 2 else fy_val
                
    return year_map

def detect_columns(df: pd.DataFrame) -> Dict[str, str]:
    mapping = {}
    cols = [str(c).strip() for c in df.columns]
    
    for c in cols:
        clow = c.lower()
        if not mapping.get("account_name") and any(k in clow for k in ["account", "particulars", "description", "name", "ledger", "item", "line item", "report date", "narration", "title", "label"]):
            mapping["account_name"] = c
        elif not mapping.get("debit") and any(k in clow for k in ["debit", "dr", "dr."]):
            mapping["debit"] = c
        elif not mapping.get("credit") and any(k in clow for k in ["credit", "cr", "cr."]):
            mapping["credit"] = c
        elif not mapping.get("amount") and any(k in clow for k in ["amount", "balance", "total", "val", "value"]):
            mapping["amount"] = c
        elif not mapping.get("account_code") and any(k in clow for k in ["code", "id", "account no", "acct#", "no."]):
            mapping["account_code"] = c
        elif not mapping.get("type") and any(k in clow for k in ["type", "group", "category", "class"]):
            mapping["type"] = c

    # Fallback if text columns are present: prioritize Column 0 if it contains text labels
    if not mapping.get("account_name"):
        if len(cols) > 0:
            col0_text_count = sum(1 for v in df[cols[0]].dropna() if isinstance(v, str) and not v.replace(".", "", 1).replace("-", "", 1).isdigit())
            if col0_text_count > 0:
                mapping["account_name"] = cols[0]
        
        if not mapping.get("account_name"):
            for c in cols:
                if df[c].dtype == "object":
                    mapping["account_name"] = c
                    break
        if not mapping.get("account_name") and len(cols) > 0:
            mapping["account_name"] = cols[0]

    # Fallback for amount column if debit, credit, and amount are all unassigned
    if not mapping.get("debit") and not mapping.get("credit") and not mapping.get("amount"):
        acct_c = mapping.get("account_name")
        for c in cols:
            if c != acct_c:
                num_count = sum(1 for v in df[c].dropna() if clean_value(v) != 0.0)
                if num_count > 0:
                    mapping["amount"] = c
                    break
        if not mapping.get("amount") and len(cols) > 1:
            mapping["amount"] = cols[1] if cols[0] == mapping.get("account_name") else cols[0]

    return mapping

import os

def parse_workbook(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    sheet_data = {}
    normalized_items = []
    detected_sheets = []
    
    fname = (filename or "").lower()
    
    meta_info = {"company_name": "Enterprise Entity", "currency": "USD", "unit": "Units"}

    try:
        if fname.endswith(".csv") or fname.endswith(".tsv"):
            df = pd.read_csv(io.BytesIO(file_bytes), sep=None, engine='python')
            sheet_data["Sheet1"] = df
            detected_sheets = ["Sheet1"]
        elif fname.endswith(".json"):
            data = json.loads(file_bytes.decode("utf-8"))
            if isinstance(data, list):
                df = pd.DataFrame(data)
                sheet_data["Sheet1"] = df
                detected_sheets = ["Sheet1"]
        elif fname.endswith(".txt"):
            lines = file_bytes.decode("utf-8").splitlines()
            parsed_lines = []
            for line in lines:
                if ":" in line:
                    parts = line.split(":", 1)
                    parsed_lines.append({
                        "account_name": parts[0].strip(),
                        "amount": parts[1].strip()
                    })
            if parsed_lines:
                df = pd.DataFrame(parsed_lines)
                sheet_data["Sheet1"] = df
                detected_sheets = ["Sheet1"]
        elif fname.endswith(".xlsx") or fname.endswith(".xls") or fname.endswith(".xlsm") or fname.endswith(".xlsb"):
            xls = pd.ExcelFile(io.BytesIO(file_bytes))
            detected_sheets = list(xls.sheet_names)
            for sheet in xls.sheet_names:
                try:
                    sheet_data[sheet] = pd.read_excel(xls, sheet_name=sheet)
                except Exception:
                    continue
        elif fname.endswith(".pdf"):
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            records = []
            
            fs_indicators = [
                "consolidated statements of income",
                "consolidated statement of income",
                "consolidated balance sheets",
                "consolidated balance sheet",
                "consolidated statements of financial position",
                "consolidated statement of financial position",
                "consolidated statements of cash flows",
                "consolidated statement of cash flows",
                "income statement",
                "balance sheet"
            ]
            
            current_years = ["2024", "2025", "2026"]
            
            for page_idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if not page_text:
                    continue
                
                page_text_lower = page_text.lower()
                if not any(ind in page_text_lower for ind in fs_indicators):
                    continue
                
                lines = page_text.splitlines()
                for line in lines:
                    line_str = line.strip()
                    if not line_str:
                        continue
                    
                    years_in_line = re.findall(r'\b(201[5-9]|202[0-9]|2030)\b', line_str)
                    if len(years_in_line) >= 2:
                        current_years = sorted(list(set(years_in_line)))
                        continue
                    
                    tokens = line_str.split()
                    if len(tokens) < 2:
                        continue
                    
                    num_tokens = []
                    text_tokens = []
                    for t in reversed(tokens):
                        t_clean = t.replace("$", "").replace("₹", "").replace("€", "").replace("£", "").replace(",", "").replace("(", "").replace(")", "").replace("[", "").replace("]", "").strip()
                        if t_clean in ["—", "–", "-", ""] or (t_clean.replace(".", "", 1).isdigit() and t_clean.replace(".", "", 1) != ""):
                            if len(t_clean) == 4 and t_clean.startswith("20") and t_clean not in current_years:
                                text_tokens.insert(0, t)
                            else:
                                num_tokens.insert(0, t)
                        else:
                            text_tokens.insert(0, t)
                            
                    if not num_tokens or not text_tokens:
                        continue
                        
                    label = " ".join(text_tokens)
                    if label.lower() in ["particulars", "notes", "notes:", "as at", "year ended", "notes as at", "description", "account name"]:
                        continue
                        
                    if len(text_tokens) > 1:
                        last_tok = text_tokens[-1]
                        if last_tok.isdigit() and int(last_tok) <= 45:
                            label = " ".join(text_tokens[:-1])
                            
                    mapped = {}
                    num_vals = [clean_value(v) for v in num_tokens]
                    
                    if len(num_vals) <= len(current_years):
                        for i, val in enumerate(num_vals):
                            yr_idx = len(current_years) - len(num_vals) + i
                            mapped[current_years[yr_idx]] = val
                    else:
                        for i in range(len(current_years)):
                            mapped[current_years[i]] = num_vals[i]
                            
                    records.append({
                        "Particulars": label,
                        **mapped
                    })
            
            if records:
                df = pd.DataFrame(records)
                sheet_data["Sheet1"] = df
                detected_sheets = ["Sheet1"]
        elif fname.endswith(".docx"):
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            records = []
            current_years = ["2024", "2025", "2026"]
            
            text_lines = []
            for p in doc.paragraphs:
                if p.text.strip():
                    text_lines.append(p.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = "  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_lines.append(row_text)
            
            for line in text_lines:
                line_str = line.strip()
                if not line_str:
                    continue
                
                years_in_line = re.findall(r'\b(201[5-9]|202[0-9]|2030)\b', line_str)
                if len(years_in_line) >= 2:
                    current_years = sorted(list(set(years_in_line)))
                    continue
                
                tokens = line_str.split()
                if len(tokens) < 2:
                    continue
                
                num_tokens = []
                text_tokens = []
                for t in reversed(tokens):
                    t_clean = t.replace("$", "").replace("₹", "").replace("€", "").replace("£", "").replace(",", "").replace("(", "").replace(")", "").replace("[", "").replace("]", "").strip()
                    if t_clean in ["—", "–", "-", ""] or (t_clean.replace(".", "", 1).isdigit() and t_clean.replace(".", "", 1) != ""):
                        if len(t_clean) == 4 and t_clean.startswith("20") and t_clean not in current_years:
                            text_tokens.insert(0, t)
                        else:
                            num_tokens.insert(0, t)
                    else:
                        text_tokens.insert(0, t)
                        
                if not num_tokens or not text_tokens:
                    continue
                    
                label = " ".join(text_tokens)
                if label.lower() in ["particulars", "notes", "notes:", "as at", "year ended", "notes as at", "description", "account name"]:
                    continue
                    
                if len(text_tokens) > 1:
                    last_tok = text_tokens[-1]
                    if last_tok.isdigit() and int(last_tok) <= 45:
                        label = " ".join(text_tokens[:-1])
                        
                mapped = {}
                num_vals = [clean_value(v) for v in num_tokens]
                
                if len(num_vals) <= len(current_years):
                    for i, val in enumerate(num_vals):
                        yr_idx = len(current_years) - len(num_vals) + i
                        mapped[current_years[yr_idx]] = val
                else:
                    for i in range(len(current_years)):
                        mapped[current_years[i]] = num_vals[i]
                        
                records.append({
                    "Particulars": label,
                    **mapped
                })
                
            if records:
                df = pd.DataFrame(records)
                sheet_data["Sheet1"] = df
                detected_sheets = ["Sheet1"]
    except Exception as e:
        print(f"Excel read error for {filename}: {e}")

    SKIP_SHEETS = ["customization", "parameters", "template", "setup", "config", "instruction", "instructions", "notes", "readme"]
    has_master_data_sheet = any("data sheet" in str(s).lower() for s in sheet_data.keys())
    if has_master_data_sheet:
        SKIP_SHEETS.extend(["profit & loss", "quarters", "balance sheet", "cash flow"])

    if sheet_data:
        meta_info = detect_company_and_currency(sheet_data, filename)

        for sheet_name, df in sheet_data.items():
            if any(skip_kw in str(sheet_name).lower().strip() for skip_kw in SKIP_SHEETS):
                continue
            if df.empty:
                continue
            df_clean = df.dropna(how="all").dropna(axis=1, how="all")
            if df_clean.empty:
                continue

            # Header detection
            header_row_idx = 0
            found_header = False
            df_body = df_clean.copy()
            
            # Check if columns are already set properly
            cols_str = " ".join([str(c).lower() for c in df_clean.columns])
            if any(kw in cols_str for kw in ["account", "particulars", "debit", "credit", "amount", "description", "balance", "item", "line item"]) and not all("unnamed" in str(c).lower() or isinstance(c, int) for c in df_clean.columns):
                found_header = True

            if not found_header:
                for r_idx in range(min(100, len(df_clean))):
                    row_vals = [str(v).lower() for v in df_clean.iloc[r_idx].values if pd.notna(v)]
                    row_str = " ".join(row_vals)
                    
                    # If row has multiple year values, it is a date header row
                    years_in_row = [v for v in row_vals if re.search(r'\b(201[5-9]|202[0-9]|2030)\b', v)]
                    if len(years_in_row) >= 2:
                        header_row_idx = r_idx
                        found_header = True
                        break

                    # Avoid matching data rows (which usually contain numbers)
                    numeric_count = sum(1 for v in df_clean.iloc[r_idx].values if pd.notna(v) and isinstance(v, (int, float)))
                    if numeric_count > 1 and len(row_vals) > 2:
                        continue
                    
                    if any(kw in row_str for kw in ["account", "particulars", "debit", "credit", "amount", "description", "balance", "item", "line item"]) or any(re.search(r'\b202[0-9]\b', v) for v in row_vals):
                        header_row_idx = r_idx
                        found_header = True
                        break

                if found_header and header_row_idx > 0:
                    header_series = df_clean.iloc[header_row_idx]
                    df_body = df_clean.iloc[header_row_idx + 1:].copy()
                    df_body.columns = [str(c).strip() if pd.notna(c) else f"Col{c_i}" for c_i, c in enumerate(header_series)]
                else:
                    df_body = df_clean.copy()

            # Multi-Year Column Detection across full sheet slice
            headers_list = list(df_body.columns)
            top_rows_list = df_clean.iloc[:100].values.tolist()
            year_col_map = detect_year_columns(headers_list, top_rows_list)

            col_map = detect_columns(df_body)
            acct_col = col_map.get("account_name")

            if not acct_col or acct_col not in df_body.columns:
                continue

            # Determine account column index
            acct_col_idx = list(df_body.columns).index(acct_col)
            current_section = ""

            for idx, row in df_body.iterrows():
                acct_name = str(row[acct_col]).strip() if pd.notna(row[acct_col]) else ""
                if not acct_name or acct_name.lower() in ["total", "subtotal", "grand total", "nan", "particulars", "account name", "description"]:
                    continue
                
                if acct_name.upper() in ["PROFIT & LOSS", "QUARTERS", "BALANCE SHEET", "CASH FLOW", "PRICE", "DERIVED"]:
                    current_section = acct_name.upper()
                    continue

                is_quarterly_item = (current_section == "QUARTERS") or ("quarters" in sheet_name.lower())
                acct_type = classify_account(acct_name, sheet_name)

                if year_col_map:
                    # Multi-Year Grid Processing: Extract value for each detected year column
                    for col_idx, year in year_col_map.items():
                        if col_idx < len(row):
                            val_raw = row.iloc[col_idx]
                            if is_blank_value(val_raw):
                                continue
                            val = clean_value(val_raw)
                            col_letter = chr(65 + col_idx) if col_idx < 26 else f"Col{col_idx}"
                            
                            try:
                                yr_val = int(year)
                            except ValueError:
                                yr_val = 2026

                            if is_quarterly_item:
                                cal_yr = yr_val
                                f_yr = "FY2027" if year == "2026" else f"FY{yr_val + 1}"
                                p_id = "Q1_FY2027" if year == "2026" else f"Q1_FY{yr_val + 1}"
                                p_start = f"{yr_val}-04-01"
                                p_end = f"{yr_val}-06-30"
                            else:
                                cal_yr = yr_val
                                f_yr = f"FY{yr_val}"
                                p_id = f"FY{yr_val}"
                                p_start = f"{yr_val - 1}-04-01"
                                p_end = f"{yr_val}-03-31"

                            if current_section == "BALANCE SHEET" or acct_type in ["ASSET", "LIABILITY", "EQUITY"] or "ASSET" in acct_type or "LIABILITY" in acct_type or "EQUITY" in acct_type:
                                stmt_type = "BALANCE_SHEET"
                            elif current_section == "CASH FLOW" or acct_type == "CASH_FLOW":
                                stmt_type = "CASH_FLOW"
                            else:
                                stmt_type = "INCOME_STATEMENT"

                            normalized_items.append({
                                "account_code": f"{sheet_name}-{idx}-{year}",
                                "account_name": acct_name,
                                "account_type": acct_type,
                                "is_summary": is_summary_or_total_row(acct_name),
                                "is_quarterly": is_quarterly_item,
                                "calendar_year": cal_yr,
                                "fiscal_year": f_yr,
                                "period_type": "QUARTERLY" if is_quarterly_item else "ANNUAL",
                                "period_id": p_id,
                                "period_start": p_start,
                                "period_end": p_end,
                                "statement_type": stmt_type,
                                "scope": "QUARTERLY" if is_quarterly_item else ("CONSOLIDATED" if "consolidated" in sheet_name.lower() or "consolidated" in acct_name.lower() else "STANDALONE"),
                                "debit": abs(val) if (acct_type in ["ASSET", "CASH_ASSET", "RECEIVABLE_ASSET", "INVENTORY_ASSET", "EXPENSE", "COGS", "DEPRECIATION_EXPENSE", "INTEREST_EXPENSE", "TAX_EXPENSE"] or "ASSET" in acct_type or "EXPENSE" in acct_type) else 0.0,
                                "credit": abs(val) if (acct_type in ["LIABILITY", "PAYABLE_LIABILITY", "DEBT_LIABILITY", "EQUITY", "REVENUE", "OPERATING_INCOME", "NET_INCOME", "SALES", "OTHER_INCOME"] or "LIABILITY" in acct_type or "REVENUE" in acct_type or "EQUITY" in acct_type) else 0.0,
                                "net_amount": val,
                                "sheet": sheet_name,
                                "row": int(idx) + 1 if isinstance(idx, (int, float)) else 1,
                                "column": col_letter,
                                "year": year,
                                "unit": meta_info["unit"],
                                "currency": meta_info["currency"],
                                "source_label": acct_name,
                                "source_value": val
                            })
                else:
                    # Single-Year / Trial Balance Fallback
                    dr_col = col_map.get("debit")
                    cr_col = col_map.get("credit")
                    amt_col = col_map.get("amount")
                    code_col = col_map.get("account_code")
                    type_col = col_map.get("type")

                    dr = clean_value(row[dr_col]) if dr_col and dr_col in row else 0.0
                    cr = clean_value(row[cr_col]) if cr_col and cr_col in row else 0.0
                    amt = clean_value(row[amt_col]) if amt_col and amt_col in row else (dr - cr)
                    
                    if dr == 0.0 and cr == 0.0 and amt == 0.0:
                        for cell_k, cell_v in row.items():
                            if str(cell_k) != acct_col and pd.notna(cell_v):
                                cv = clean_value(cell_v)
                                if cv != 0.0:
                                    amt = cv
                                    break

                    acct_code = str(row[code_col]).strip() if code_col and code_col in row and pd.notna(row[code_col]) else f"ACC-{idx}"
                    raw_type = str(row[type_col]).strip() if type_col and type_col in row and pd.notna(row[type_col]) else ""
                    if raw_type:
                        acct_type = raw_type.upper()
                    net = dr - cr if (dr or cr) else amt

                    normalized_items.append({
                        "account_code": acct_code,
                        "account_name": acct_name,
                        "account_type": acct_type,
                        "is_summary": is_summary_or_total_row(acct_name),
                        "debit": clean_value(dr),
                        "credit": clean_value(cr),
                        "net_amount": clean_value(net),
                        "sheet": sheet_name,
                        "row": int(idx) + 1 if isinstance(idx, (int, float)) else 1,
                        "column": "A",
                        "year": "Current",
                        "unit": meta_info["unit"],
                        "currency": meta_info["currency"],
                        "source_label": acct_name,
                        "source_value": clean_value(net)
                    })

    # Collect list of detected years
    all_years = sorted(list(set(i["year"] for i in normalized_items if i.get("year"))))
    if not all_years:
        all_years = ["Current"]

    if not detected_sheets:
        detected_sheets = sorted(list(set(item["sheet"] for item in normalized_items))) if normalized_items else ["ParsedSheet"]

    return {
        "filename": filename,
        "company_name": meta_info["company_name"],
        "currency": meta_info["currency"],
        "unit": meta_info["unit"],
        "sheet_names": detected_sheets,
        "years": all_years,
        "normalized_items": normalized_items
    }
